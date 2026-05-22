from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
import time
import uuid
from datetime import UTC, datetime
from typing import Any

import sqlparse
from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import Response, StreamingResponse
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Font
from pydantic import BaseModel, Field
from sqlalchemy import func as sa_func, text as sa_text
from sqlalchemy.orm import Session

from lib.database import get_database
from lib.database import SessionLocal as AppSessionLocal
from lib.megabase import SessionLocal as MegabaseSessionLocal
from lib.megabase import drop_table as megabase_drop_table
from lib.megabase import init_megabase
from lib.megabase import describe_table as megabase_describe_table
from lib.megabase import query_records as megabase_query_records
from lib.ai import get_generative_model
from lib.models import Asset, LogGroup, LogFile, LogMessage, LogProcess, LogReport, LogTable, LogTableSummary, User
from lib.storage import delete_file, download_file, upload_file
from parsers.extra_grouping import group_rows_by_extra
from parsers.orchestrator import create_process, enqueue_process, mark_process_failed
from routes.auth import get_current_user

router = APIRouter(prefix="/logs", tags=["logs"])
logger = logging.getLogger(__name__)


class MessageResponse(BaseModel):
    message: str


class CreateLogGroupRequest(BaseModel):
    name: str
    profile_name: str | None = "default"


class UpdateLogGroupRequest(BaseModel):
    name: str
    profile_name: str | None = None


class CreateProcessRequest(BaseModel):
    file_ids: list[str] | None = None


class LogGroupResponse(BaseModel):
    id: str
    user_id: str
    name: str
    profile_name: str | None
    created_at: datetime
    file_count: int = 0
    table_count: int = 0


class LogFileResponse(BaseModel):
    id: str
    group_id: str
    asset_id: str
    name: str
    size: int
    content_type: str
    created_at: datetime


class LogProcessResponse(BaseModel):
    id: str
    group_id: str
    file_id: str | None
    status: str
    classification: dict[str, Any] | None
    result: dict[str, Any] | None
    error: str | None
    created_at: datetime
    updated_at: datetime


class FileProcessOutcomeResponse(BaseModel):
    file_id: str
    filename: str
    process_id: str | None
    status: str
    error: str | None = None


class UploadFilesResponse(BaseModel):
    process_ids: list[str]
    status: str
    files: list[LogFileResponse]
    outcomes: list[FileProcessOutcomeResponse]


class ProcessEnqueuedResponse(BaseModel):
    process_ids: list[str]
    status: str
    errors: list[str] = Field(default_factory=list)


class PersistedMessagesResponse(BaseModel):
    messages: list[dict[str, Any]]


class ReplaceMessagesRequest(BaseModel):
    messages: list[dict[str, Any]]


class ReplaceMessagesResponse(BaseModel):
    saved_messages: int


class ChatRequest(BaseModel):
    message: str
    history: list[dict[str, Any]] = Field(default_factory=list)


class LogInsightReport(BaseModel):
    summary: str
    severity: str
    top_errors: list[str] = Field(default_factory=list)
    root_cause_hypothesis: str
    log_sequence_narrative: str
    recommendations: list[str] = Field(default_factory=list)
    anomalies: list[str] = Field(default_factory=list)


class TableSummaryResponse(BaseModel):
    summary: str
    key_observations: list[str] = Field(default_factory=list)
    severity: str
    next_actions: list[str] = Field(default_factory=list)
    errors_or_anomalies: list[str] = Field(default_factory=list)


class QueryRequest(BaseModel):
    sql: str


class FilteredExportRequest(BaseModel):
    format: str = "csv"
    search: str | None = None
    levels: list[str] = Field(default_factory=list)
    field_filters: dict[str, str] = Field(default_factory=dict)
    timestamp_from: str | None = None
    timestamp_to: str | None = None


class QueryResponse(BaseModel):
    status: str
    columns: list[str] = Field(default_factory=list)
    rows: list[list[Any]] = Field(default_factory=list)
    row_count: int = 0
    execution_time_ms: float = 0
    message: str = ""


class ReportSectionTable(BaseModel):
    title: str
    columns: list[str]
    rows: list[list[Any]]


class ReportSectionRequest(BaseModel):
    heading: str
    content: str
    tables: list[ReportSectionTable] = Field(default_factory=list)


class ReportRequest(BaseModel):
    title: str
    sections: list[ReportSectionRequest]


FORBIDDEN_SQL_KEYWORDS = re.compile(
    r"^\s*(INSERT|UPDATE|DELETE|DROP|ALTER|CREATE|TRUNCATE|REPLACE|GRANT|REVOKE)\s",
    re.IGNORECASE,
)

# Maximum LIMIT that may be injected into any SQL query for safety.
_HARD_SQL_LIMIT = 10000

# Chat message limits.
_MAX_PERSISTED_MESSAGES = 500
_MAX_MESSAGE_CONTENT_LENGTH = 100000


def _extract_table_names(parsed: sqlparse.sql.Statement) -> set[str]:
    """Extract table names referenced in a parsed SQL statement."""
    tables: set[str] = set()

    def extract_from_token(token: sqlparse.sql.Token) -> None:
        if isinstance(token, sqlparse.sql.Identifier):
            name = token.get_real_name()
            if name:
                tables.add(name)
        elif isinstance(token, sqlparse.sql.IdentifierList):
            for identifier in token.get_identifiers():
                name = identifier.get_real_name()
                if name:
                    tables.add(name)
        elif token.ttype is sqlparse.tokens.Name:
            tables.add(token.value)
        elif isinstance(token, sqlparse.sql.Where):
            # Stop traversal at WHERE clause - table names can't appear there
            return

    if parsed.get_type() == "SELECT":
        from_seen = False
        for token in parsed.tokens:
            if token.ttype is sqlparse.tokens.Keyword and token.value.upper() in {
                "FROM",
                "JOIN",
                "INNER JOIN",
                "LEFT JOIN",
                "RIGHT JOIN",
                "FULL JOIN",
                "LEFT OUTER JOIN",
                "RIGHT OUTER JOIN",
                "FULL OUTER JOIN",
                "CROSS JOIN",
                "NATURAL JOIN",
            }:
                from_seen = True
                continue
            if from_seen:
                extract_from_token(token)
                from_seen = False
            if token.ttype is sqlparse.tokens.Keyword and token.value.upper() in {
                "WHERE",
                "GROUP",
                "ORDER",
                "HAVING",
                "LIMIT",
                "OFFSET",
            }:
                break
    return tables


def _validate_table_allowlist(sql_text: str, allowed_tables: set[str]) -> str | None:
    """Validate that all tables referenced in the SQL are in the allowed set.

    Returns an error message string if validation fails, or None on success.
    """
    parsed = sqlparse.parse(sql_text)
    for statement in parsed:
        refs = _extract_table_names(statement)
        if not refs:
            continue
        disallowed = refs - allowed_tables
        if disallowed:
            return f"Query references disallowed tables: {', '.join(sorted(disallowed))}."

    return None


def _inject_sql_limit(sql_text: str, max_rows: int = _HARD_SQL_LIMIT) -> str:
    """Inject a LIMIT clause into a SELECT if one is not already present."""
    parsed = sqlparse.parse(sql_text)
    if not parsed:
        return sql_text

    stmt = parsed[0]
    if stmt.get_type() != "SELECT":
        return sql_text

    # Check if LIMIT is already present
    has_limit = any(token.ttype is sqlparse.tokens.Keyword and token.value.upper() == "LIMIT" for token in stmt.tokens)
    if has_limit:
        return sql_text

    return f"{sql_text.rstrip(';').rstrip()} LIMIT {max_rows};"


QUERY_RESULT_LIMIT = 500


def _uuid_or_raw(value: str):
    try:
        return uuid.UUID(str(value))
    except ValueError:
        return value


def _parse_json(value: str | None):
    if not value:
        return None

    try:
        parsed = json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return {"raw": value}

    if isinstance(parsed, dict):
        return parsed

    return {"value": parsed}


def _parse_message_payload(payload: str | None, role: str, content: str):
    if payload is None or payload == "":
        return {"role": role, "content": content}

    try:
        parsed = json.loads(payload)
    except (json.JSONDecodeError, TypeError):
        return {"role": role, "content": content}

    if isinstance(parsed, dict):
        return parsed

    return {"role": role, "content": content}


def _group_response(group: LogGroup, file_count: int = 0, table_count: int = 0):
    return LogGroupResponse(
        id=str(group.id),
        user_id=str(group.user_id),
        name=group.name,
        profile_name=group.profile_name,
        created_at=group.created_at,
        file_count=file_count,
        table_count=table_count,
    )


def _log_file_response(log_file: LogFile, asset: Asset):
    return LogFileResponse(
        id=str(log_file.id),
        group_id=str(log_file.group_id),
        asset_id=str(log_file.asset_id),
        name=asset.name,
        size=asset.size,
        content_type=asset.type,
        created_at=log_file.created_at,
    )


def _log_process_response(process: LogProcess):
    return LogProcessResponse(
        id=str(process.id),
        group_id=str(process.group_id),
        file_id=str(process.file_id) if process.file_id is not None else None,
        status=process.status,
        classification=_parse_json(process.classification),
        result=_parse_json(process.result),
        error=process.error,
        created_at=process.created_at,
        updated_at=process.updated_at,
    )


def _batched_status(total: int, queued: int) -> str:
    if queued == 0:
        return "failed"
    if queued == total:
        return "queued"
    return "partial"


def _require_owned_group(database: Session, group_id: str, user_id: uuid.UUID):
    group = database.query(LogGroup).filter(LogGroup.id == _uuid_or_raw(group_id), LogGroup.user_id == user_id).first()
    if group is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Log group not found.")
    return group


def _require_owned_file(database: Session, group_id: str, file_id: str):
    log_file = (
        database.query(LogFile)
        .filter(LogFile.id == _uuid_or_raw(file_id), LogFile.group_id == _uuid_or_raw(group_id))
        .first()
    )
    if log_file is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Log file not found.")

    asset = database.query(Asset).filter(Asset.id == log_file.asset_id).first()
    if asset is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Asset not found.")

    return log_file, asset


def _delete_orphan_assets(asset_ids: list[uuid.UUID]):
    database = AppSessionLocal()
    try:
        for asset_id in asset_ids:
            try:
                remaining_links = database.query(LogFile).filter(LogFile.asset_id == asset_id).count()
                if remaining_links == 0:
                    delete_file(asset_id=asset_id, db=database)
            except Exception:  # noqa: BLE001
                logger.exception("Failed to delete orphan asset %s", asset_id)
    finally:
        database.close()


def _cleanup_generated_tables_for_file(database: Session, group_id: str, file_id: str) -> None:
    latest_completed_process = (
        database.query(LogProcess)
        .filter(
            LogProcess.group_id == _uuid_or_raw(group_id),
            LogProcess.file_id == _uuid_or_raw(file_id),
            LogProcess.status == "completed",
        )
        .order_by(LogProcess.updated_at.desc())
        .first()
    )

    if latest_completed_process is None or not latest_completed_process.result:
        return

    parsed_result = _parse_json(latest_completed_process.result)
    if not isinstance(parsed_result, dict):
        return

    table_definitions = parsed_result.get("table_definitions")
    if not isinstance(table_definitions, list):
        return

    table_names: set[str] = set()
    for table_definition in table_definitions:
        if isinstance(table_definition, dict):
            table_name = table_definition.get("table_name")
            if isinstance(table_name, str) and table_name:
                table_names.add(table_name)

    if not table_names:
        return

    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        for table_name in table_names:
            try:
                megabase_drop_table(megabase_database, table_name)
            except Exception:  # noqa: BLE001
                logger.exception("Failed to drop generated table %s before reprocessing file %s", table_name, file_id)
    finally:
        megabase_database.close()

    (
        database.query(LogTable)
        .filter(LogTable.group_id == _uuid_or_raw(group_id), LogTable.table.in_(table_names))
        .delete(synchronize_session=False)
    )
    database.commit()


@router.get("", response_model=list[LogGroupResponse])
def list_log_groups(
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    file_count_subq = (
        database.query(LogFile.group_id, sa_func.count(LogFile.id).label("file_count"))
        .group_by(LogFile.group_id)
        .subquery()
    )
    table_count_subq = (
        database.query(LogTable.group_id, sa_func.count(LogTable.id).label("table_count"))
        .group_by(LogTable.group_id)
        .subquery()
    )
    entries = (
        database.query(
            LogGroup,
            sa_func.coalesce(file_count_subq.c.file_count, 0),
            sa_func.coalesce(table_count_subq.c.table_count, 0),
        )
        .outerjoin(file_count_subq, LogGroup.id == file_count_subq.c.group_id)
        .outerjoin(table_count_subq, LogGroup.id == table_count_subq.c.group_id)
        .filter(LogGroup.user_id == current_user.id)
        .order_by(LogGroup.created_at.desc())
        .all()
    )
    return [_group_response(group, file_count=fcount, table_count=tcount) for group, fcount, tcount in entries]


@router.post("", response_model=LogGroupResponse, status_code=status.HTTP_201_CREATED)
def create_log_group(
    payload: CreateLogGroupRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Group name must not be empty.")

    profile_name = (payload.profile_name or "default").strip() or "default"

    group = LogGroup(user_id=current_user.id, name=name, profile_name=profile_name)
    database.add(group)
    database.commit()
    database.refresh(group)
    return _group_response(group)


@router.get("/{group_id}", response_model=LogGroupResponse)
def get_log_group(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    file_count = database.query(sa_func.count(LogFile.id)).filter(LogFile.group_id == group.id).scalar()
    table_count = database.query(sa_func.count(LogTable.id)).filter(LogTable.group_id == group.id).scalar()
    return _group_response(group, file_count=file_count, table_count=table_count)


@router.patch("/{group_id}", response_model=LogGroupResponse)
def update_log_group(
    group_id: str,
    payload: UpdateLogGroupRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Group name must not be empty.")

    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    group.name = name
    if payload.profile_name is not None:
        group.profile_name = payload.profile_name.strip() or "default"
    database.commit()
    database.refresh(group)
    return _group_response(group)


@router.delete("/{group_id}", response_model=MessageResponse)
def delete_log_group(
    group_id: str,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    # 1. Collect megabase table names and orphan asset IDs before cascade deletion
    log_tables = database.query(LogTable).filter(LogTable.group_id == group.id).all()
    table_names = [table.table for table in log_tables]

    file_rows = database.query(LogFile).filter(LogFile.group_id == group.id).all()
    orphan_asset_ids = {file_row.asset_id for file_row in file_rows}

    # 2. Drop all megabase tables associated with this group
    if table_names:
        megabase_database = MegabaseSessionLocal()
        try:
            init_megabase(megabase_database)
            for table_name in table_names:
                try:
                    megabase_drop_table(megabase_database, table_name)
                except Exception:  # noqa: BLE001
                    logger.exception("Failed to drop generated table %s", table_name)
        finally:
            megabase_database.close()

    # 3. Delete the group — ORM cascades to all children (files, tables, messages, processes, reports)
    database.delete(group)
    database.commit()

    # 4. Delete orphan S3 assets (no remaining LogFile references)
    if orphan_asset_ids:
        background_tasks.add_task(_delete_orphan_assets, list(orphan_asset_ids))

    return MessageResponse(message="Log group deleted.")


@router.get("/{group_id}/files", response_model=list[LogFileResponse])
def list_log_files(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    rows = database.query(LogFile).filter(LogFile.group_id == group.id).order_by(LogFile.created_at.desc()).all()

    responses: list[LogFileResponse] = []
    for row in rows:
        asset = database.query(Asset).filter(Asset.id == row.asset_id).first()
        if asset is None:
            continue
        responses.append(_log_file_response(row, asset))

    return responses


@router.get("/{group_id}/files/{file_id}", response_model=LogFileResponse)
def get_log_file_metadata(
    group_id: str,
    file_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    log_file, asset = _require_owned_file(database=database, group_id=str(group.id), file_id=file_id)
    return _log_file_response(log_file, asset)


@router.get("/{group_id}/files/{file_id}/download")
def download_log_file(
    group_id: str,
    file_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    log_file, asset = _require_owned_file(database=database, group_id=str(group.id), file_id=file_id)

    payload = download_file(asset_id=log_file.asset_id, db=database)
    if payload is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File content not found.")

    headers = {"Content-Disposition": f'attachment; filename="{asset.name}"'}
    return Response(content=payload, media_type=asset.type or "application/octet-stream", headers=headers)


def _get_table_records(group_id: str, table_name: str, current_user: User, database: Session) -> list[dict]:
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    table_record = database.query(LogTable).filter(LogTable.group_id == group.id, LogTable.table == table_name).first()
    if table_record is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Table not found.")

    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        records = megabase_query_records(megabase_database, table_name, limit=100000)
    finally:
        megabase_database.close()

    return group_rows_by_extra(records)


def _serialize_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=True)
    return str(value)


def _serialize_record_value(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=True)
    return value


def _normalize_level(value: Any) -> str:
    if value is None:
        return ""
    normalized = str(value).strip().upper()
    if normalized == "WARNING":
        return "WARN"
    return normalized


def _extract_row_level(record: dict[str, Any]) -> str:
    for key in ("log_level", "level", "severity"):
        if key in record:
            return _normalize_level(record.get(key))
    return ""


def _extract_row_timestamp(record: dict[str, Any]) -> datetime | None:
    for key in ("timestamp", "ts", "time"):
        value = record.get(key)
        if value is None:
            continue
        text = str(value).strip()
        if not text:
            continue
        normalized = text.replace("Z", "+00:00")
        try:
            parsed = datetime.fromisoformat(normalized)
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=UTC)
            return parsed.astimezone(UTC)
        except ValueError:
            continue
    return None


def _coerce_datetime(raw_value: str | None) -> datetime | None:
    if raw_value is None:
        return None
    normalized = raw_value.strip().replace("Z", "+00:00")
    if normalized == "":
        return None
    try:
        parsed = datetime.fromisoformat(normalized)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=UTC)
        return parsed.astimezone(UTC)
    except ValueError:
        return None


def _record_matches_filters(record: dict[str, Any], payload: FilteredExportRequest) -> bool:
    if payload.search:
        search_text = payload.search.strip().lower()
        if search_text:
            haystack = json.dumps(record, ensure_ascii=True, sort_keys=True).lower()
            if search_text not in haystack:
                return False

    if payload.levels:
        allowed_levels = {_normalize_level(level) for level in payload.levels if level.strip()}
        if allowed_levels:
            record_level = _extract_row_level(record)
            if record_level not in allowed_levels:
                return False

    if payload.field_filters:
        for key, expected in payload.field_filters.items():
            expected_text = expected.strip().lower()
            if expected_text == "":
                continue
            actual = record.get(key)
            actual_text = "" if actual is None else str(actual).lower()
            if expected_text not in actual_text:
                return False

    from_dt = _coerce_datetime(payload.timestamp_from)
    to_dt = _coerce_datetime(payload.timestamp_to)
    if from_dt is not None or to_dt is not None:
        row_timestamp = _extract_row_timestamp(record)
        if row_timestamp is None:
            return False
        if from_dt is not None and row_timestamp < from_dt:
            return False
        if to_dt is not None and row_timestamp > to_dt:
            return False

    return True


def _apply_export_filters(records: list[dict[str, Any]], payload: FilteredExportRequest) -> list[dict[str, Any]]:
    return [record for record in records if _record_matches_filters(record, payload)]


@router.get("/{group_id}/tables/{table_name}/download/csv")
def download_table_csv(
    group_id: str,
    table_name: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    records = _get_table_records(group_id, table_name, current_user, database)

    if not records:
        columns: list[str] = []
    else:
        columns = list(records[0].keys())

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(columns)
    for record in records:
        writer.writerow([_serialize_value(record.get(col)) for col in columns])

    content = output.getvalue()
    filename = f"{table_name}.csv"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(content=content, media_type="text/csv", headers=headers)


@router.get("/{group_id}/tables/{table_name}/download/xlsx")
def download_table_xlsx(
    group_id: str,
    table_name: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    records = _get_table_records(group_id, table_name, current_user, database)

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = table_name[:31]

    if not records:
        columns: list[str] = []
    else:
        columns = list(records[0].keys())

    worksheet.append(columns)
    for record in records:
        row_values = [_serialize_record_value(record.get(col)) for col in columns]
        worksheet.append(row_values)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)

    filename = f"{table_name}.xlsx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )


@router.post("/{group_id}/tables/{table_name}/download/filtered")
def download_table_filtered(
    group_id: str,
    table_name: str,
    payload: FilteredExportRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    records = _get_table_records(group_id, table_name, current_user, database)
    filtered_records = _apply_export_filters(records, payload)

    export_format = payload.format.strip().lower()
    if export_format not in {"csv", "json"}:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="format must be csv or json")

    if export_format == "json":
        serialized = [{k: _serialize_record_value(v) for k, v in record.items()} for record in filtered_records]
        content = json.dumps(serialized, ensure_ascii=True)
        headers = {"Content-Disposition": f'attachment; filename="{table_name}.filtered.json"'}
        return Response(content=content, media_type="application/json", headers=headers)

    columns: list[str] = []
    if filtered_records:
        seen_columns: dict[str, None] = {}
        for row in filtered_records:
            for key in row.keys():
                seen_columns.setdefault(key, None)
        columns = list(seen_columns.keys())

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(columns)
    for record in filtered_records:
        writer.writerow([_serialize_value(record.get(column)) for column in columns])

    headers = {"Content-Disposition": f'attachment; filename="{table_name}.filtered.csv"'}
    return Response(content=output.getvalue(), media_type="text/csv", headers=headers)


@router.delete("/{group_id}/files/{file_id}", response_model=MessageResponse)
def delete_log_file_route(
    group_id: str,
    file_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    log_file, _ = _require_owned_file(database=database, group_id=str(group.id), file_id=file_id)

    # 1. Drop megabase tables created by this file's processes
    _cleanup_generated_tables_for_file(database=database, group_id=str(group.id), file_id=file_id)

    # 2. Delete the file — cascades to associated processes
    asset_id = log_file.asset_id
    database.delete(log_file)
    database.commit()

    # 3. Delete the underlying asset if no other files reference it
    remaining_links = database.query(LogFile).filter(LogFile.asset_id == asset_id).count()
    if remaining_links == 0:
        delete_file(asset_id=asset_id, db=database)

    return MessageResponse(message="Log file deleted.")


@router.post("/{group_id}/files/upload", response_model=UploadFilesResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_log_files(
    group_id: str,
    files: list[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    if not files:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="At least one file is required.")

    uploaded_files: list[LogFileResponse] = []
    process_ids: list[str] = []
    outcomes: list[FileProcessOutcomeResponse] = []

    for file in files:
        filename = (file.filename or "uploaded.log").strip() or "uploaded.log"
        content_type = file.content_type or "application/octet-stream"
        file_data = await file.read()

        asset = upload_file(file_data=file_data, filename=filename, content_type=content_type, db=database)

        log_file = LogFile(
            user_id=current_user.id,
            asset_id=asset.id,
            group_id=group.id,
        )
        database.add(log_file)
        database.commit()
        database.refresh(log_file)

        file_id = str(log_file.id)
        uploaded_files.append(_log_file_response(log_file, asset))

        process_id: str | None = None
        try:
            process_id = create_process(
                group_id=str(group.id),
                file_ids=[file_id],
                file_id=file_id,
            )
            enqueue_process(
                process_id=process_id,
                group_id=str(group.id),
                file_ids_json=json.dumps([file_id], ensure_ascii=True),
            )
            process_ids.append(process_id)
            outcomes.append(
                FileProcessOutcomeResponse(
                    file_id=file_id,
                    filename=asset.name,
                    process_id=process_id,
                    status="queued",
                )
            )
        except Exception as error:  # noqa: BLE001
            logger.exception("Failed to enqueue process for file %s", file_id)
            if process_id is not None:
                mark_process_failed(
                    process_id=process_id,
                    group_id=str(group.id),
                    message=f"Queueing failed: {error}",
                )
            outcomes.append(
                FileProcessOutcomeResponse(
                    file_id=file_id,
                    filename=asset.name,
                    process_id=process_id,
                    status="failed",
                    error=str(error),
                )
            )

    return UploadFilesResponse(
        process_ids=process_ids,
        status=_batched_status(total=len(outcomes), queued=len(process_ids)),
        files=uploaded_files,
        outcomes=outcomes,
    )


@router.get("/{group_id}/processes", response_model=list[LogProcessResponse])
def list_entry_processes(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    processes = (
        database.query(LogProcess).filter(LogProcess.group_id == group.id).order_by(LogProcess.created_at.desc()).all()
    )
    return [_log_process_response(process) for process in processes]


@router.get("/{group_id}/processes/{process_id}", response_model=LogProcessResponse)
def get_entry_process(
    group_id: str,
    process_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    process = (
        database.query(LogProcess)
        .filter(LogProcess.id == _uuid_or_raw(process_id), LogProcess.group_id == group.id)
        .first()
    )
    if process is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Process not found.")
    return _log_process_response(process)


@router.post("/{group_id}/processes", response_model=ProcessEnqueuedResponse, status_code=status.HTTP_202_ACCEPTED)
def create_entry_process(
    group_id: str,
    payload: CreateProcessRequest | None = None,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    selected_file_ids: list[str] = []
    if payload and payload.file_ids:
        for file_id in payload.file_ids:
            normalized_id = str(_uuid_or_raw(file_id))
            exists = (
                database.query(LogFile)
                .filter(LogFile.id == _uuid_or_raw(file_id), LogFile.group_id == group.id)
                .first()
            )
            if exists is None:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Log file '{file_id}' not found.")
            selected_file_ids.append(normalized_id)
    else:
        selected_file_ids = [
            str(file_row.id) for file_row in database.query(LogFile).filter(LogFile.group_id == group.id)
        ]

    if not selected_file_ids:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="No files available to process.")

    process_ids: list[str] = []
    errors: list[str] = []
    for file_id in selected_file_ids:
        process_id: str | None = None
        try:
            _cleanup_generated_tables_for_file(database=database, group_id=str(group.id), file_id=file_id)
            process_id = create_process(group_id=str(group.id), file_ids=[file_id], file_id=file_id)
            enqueue_process(
                process_id=process_id,
                group_id=str(group.id),
                file_ids_json=json.dumps([file_id], ensure_ascii=True),
            )
            process_ids.append(process_id)
        except Exception as error:  # noqa: BLE001
            logger.exception("Failed to enqueue process for file %s", file_id)
            if process_id is not None:
                mark_process_failed(
                    process_id=process_id,
                    group_id=str(group.id),
                    message=f"Queueing failed: {error}",
                )
            errors.append(f"{file_id}: {error}")

    return ProcessEnqueuedResponse(
        process_ids=process_ids,
        status=_batched_status(total=len(selected_file_ids), queued=len(process_ids)),
        errors=errors,
    )


@router.get("/{group_id}/chat/messages", response_model=PersistedMessagesResponse)
def get_chat_messages(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    messages = (
        database.query(LogMessage).filter(LogMessage.group_id == group.id).order_by(LogMessage.created_at.asc()).all()
    )

    return PersistedMessagesResponse(
        messages=[_parse_message_payload(message.payload, message.role, message.content) for message in messages]
    )


@router.put("/{group_id}/chat/messages", response_model=ReplaceMessagesResponse)
def replace_chat_messages(
    group_id: str,
    payload: ReplaceMessagesRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    database.query(LogMessage).filter(LogMessage.group_id == group.id).delete(synchronize_session="fetch")

    saved_messages = 0
    messages_to_save = payload.messages[:_MAX_PERSISTED_MESSAGES]
    if len(payload.messages) > _MAX_PERSISTED_MESSAGES:
        logger.warning(
            "Truncated chat messages for group %s: %d -> %d",
            group.id,
            len(payload.messages),
            _MAX_PERSISTED_MESSAGES,
        )

    for message in messages_to_save:
        role = str(message.get("role", "assistant"))
        if role == "":
            role = "assistant"

        content = message.get("content", "")
        if isinstance(content, str):
            normalized_content = content[:_MAX_MESSAGE_CONTENT_LENGTH]
        else:
            normalized_content = json.dumps(content, ensure_ascii=True)[:_MAX_MESSAGE_CONTENT_LENGTH]

        database.add(
            LogMessage(
                group_id=group.id,
                role=role,
                content=normalized_content,
                payload=json.dumps(message, ensure_ascii=True),
            )
        )
        saved_messages += 1

    database.commit()

    return ReplaceMessagesResponse(saved_messages=saved_messages)


MAX_CONTEXT_ROWS = 200


def _fetch_group_table_context(group_id: str, database: Session) -> str:
    tables = database.query(LogTable).filter(LogTable.group_id == _uuid_or_raw(group_id)).all()
    if not tables:
        return "No parsed tables are available for this log group."

    lines: list[str] = ["Available tables and sample rows:", ""]
    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        for table in tables:
            lines.append(f"Table: {table.table}")
            lines.append(f"Schema: {table.schema}")
            try:
                result = megabase_database.execute(sa_text(f'SELECT * FROM "{table.table}" LIMIT 20'))
                columns = [str(col) for col in result.keys()]
                rows = result.fetchall()
                lines.append(f"Columns: {', '.join(columns)}")
                lines.append(f"Sample rows ({len(rows)}):")
                for row in rows[:5]:
                    row_dict = dict(zip(columns, row))
                    lines.append(json.dumps(row_dict, ensure_ascii=True, default=str))
                lines.append("")
            except Exception as error:
                lines.append(f"Could not sample rows: {error}")
                lines.append("")
    finally:
        megabase_database.close()

    return "\n".join(lines)


def _build_chat_system_prompt(group_name: str) -> str:
    return (
        "You are Logdog's AI log analysis assistant. "
        f'You are helping the user analyze the log group "{group_name}".\n\n'
        "Refer to this log group by its display name. Do not mention internal IDs or UUIDs.\n"
        "Answer the user's questions based on the log data context provided in the first message below. "
        "Be concise, accurate, and actionable. If the data is insufficient, say so."
    )


@router.post("/{group_id}/chat")
def stream_chat(
    group_id: str,
    payload: ChatRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    table_context = _fetch_group_table_context(str(group.id), database)
    group_name = (group.name or "").strip().replace("\n", " ")[:120] or f"log group {group_id}"

    system_prompt = _build_chat_system_prompt(group_name)

    model = get_generative_model()
    messages: list[tuple[str, str]] = [("system", system_prompt)]

    # Inject table context as a clearly delimited user message rather than
    # embedding untrusted log content in the system prompt (prompt-injection mitigation).
    messages.append(("user", f"[Log data context for this session]\n\n{table_context}\n\n---\n"))

    for turn in payload.history:
        role = str(turn.get("role", "user"))
        content = str(turn.get("content", ""))
        if role in ("user", "assistant"):
            messages.append((role, content))
    messages.append(("user", payload.message))

    def event_generator():
        for chunk in model.client.stream(messages):
            if chunk.content:
                data = json.dumps({"token": chunk.content})
                yield f"data: {data}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


MAX_REPORT_ROWS = 200


def _fetch_group_rows_for_report(group_id: str, database: Session) -> str:
    tables = database.query(LogTable).filter(LogTable.group_id == _uuid_or_raw(group_id)).all()
    if not tables:
        return "No parsed tables are available for this log group."

    lines: list[str] = ["Log data context:", ""]
    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        for table in tables:
            lines.append(f"Table: {table.table}")
            try:
                result = megabase_database.execute(sa_text(f'SELECT * FROM "{table.table}" LIMIT {MAX_REPORT_ROWS}'))
                columns = [str(col) for col in result.keys()]
                rows = result.fetchall()
                lines.append(f"Columns: {', '.join(columns)}")
                lines.append(f"Row count in sample: {len(rows)}")
                for row in rows[:10]:
                    row_dict = dict(zip(columns, row))
                    lines.append(json.dumps(row_dict, ensure_ascii=True, default=str))
                lines.append("")
            except Exception as error:
                lines.append(f"Could not read table: {error}")
                lines.append("")
    finally:
        megabase_database.close()

    return "\n".join(lines)


class GroupStatsResponse(BaseModel):
    process_count: int
    process_status_counts: dict[str, int]
    file_count: int
    file_format_counts: dict[str, int]
    table_count: int
    table_row_counts: dict[str, int]
    total_rows: int
    parser_confidence: float | None = None


@router.get("/{group_id}/stats", response_model=GroupStatsResponse)
def get_group_stats(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    # Process counts
    group_processes = database.query(LogProcess).filter(LogProcess.group_id == group.id).all()
    process_count = len(group_processes)
    process_status_counts: dict[str, int] = {}
    confidence_sum = 0.0
    confidence_count = 0
    for proc in group_processes:
        process_status_counts[proc.status] = process_status_counts.get(proc.status, 0) + 1
        if proc.status == "completed" and proc.result:
            try:
                parsed = json.loads(proc.result)
                if isinstance(parsed, dict):
                    conf = parsed.get("confidence")
                    if isinstance(conf, (int, float)):
                        confidence_sum += max(0.0, min(float(conf), 1.0))
                        confidence_count += 1
            except (json.JSONDecodeError, TypeError):
                pass
    parser_confidence = round(confidence_sum / confidence_count, 2) if confidence_count > 0 else None

    # File counts by format
    files = database.query(LogFile).filter(LogFile.group_id == group.id).all()
    file_count = len(files)
    file_format_counts: dict[str, int] = {}
    for log_file in files:
        asset = database.query(Asset).filter(Asset.id == log_file.asset_id).first()
        if asset is not None:
            ext = os.path.splitext(asset.name)[1].lower() or "unknown"
            file_format_counts[ext] = file_format_counts.get(ext, 0) + 1

    # Table row counts
    tables = database.query(LogTable).filter(LogTable.group_id == group.id).all()
    table_count = len(tables)
    table_row_counts: dict[str, int] = {}
    total_rows = 0
    if tables:
        megabase_database = MegabaseSessionLocal()
        try:
            init_megabase(megabase_database)
            for table in tables:
                try:
                    result = megabase_database.execute(sa_text(f'SELECT COUNT(*) FROM "{table.table}"'))
                    row_count = result.scalar() or 0
                    table_row_counts[table.table] = row_count
                    total_rows += row_count
                except Exception:
                    table_row_counts[table.table] = 0
        finally:
            megabase_database.close()

    return GroupStatsResponse(
        process_count=process_count,
        process_status_counts=process_status_counts,
        file_count=file_count,
        file_format_counts=file_format_counts,
        table_count=table_count,
        table_row_counts=table_row_counts,
        total_rows=total_rows,
        parser_confidence=parser_confidence,
    )


@router.post("/{group_id}/insights", response_model=LogInsightReport)
def generate_insights(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    context = _fetch_group_rows_for_report(str(group.id), database)

    system_prompt = (
        "You are an expert log analyst. Analyze the provided log data and generate a structured insight report. "
        "Return valid JSON matching the required schema. Be concise and factual."
    )

    prompt = (
        f"{context}\n\n"
        "Generate a JSON report with these fields:\n"
        '- "summary": a 1-2 sentence overview\n'
        '- "severity": one of low, medium, high, critical\n'
        '- "top_errors": list of top error strings found\n'
        '- "root_cause_hypothesis": a brief hypothesis\n'
        '- "log_sequence_narrative": a short narrative of what happened\n'
        '- "recommendations": list of actionable recommendations\n'
        '- "anomalies": list of anomalies detected\n'
    )

    model = get_generative_model()
    report = model.generate_structured(prompt, LogInsightReport, system_prompt=system_prompt)

    log_report = LogReport(
        group_id=group.id,
        content=report.model_dump(),
    )
    database.add(log_report)
    database.commit()

    return report


@router.get("/{group_id}/insights", response_model=LogInsightReport | None)
def get_insights(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    log_report = (
        database.query(LogReport).filter(LogReport.group_id == group.id).order_by(LogReport.created_at.desc()).first()
    )

    if log_report is None:
        return None

    return LogInsightReport(**log_report.content)


SUMMARY_ROWS_LIMIT = 200
SUMMARY_CACHE_KEY = "v1"


def _fetch_table_rows_for_summary(group_id: str, table_name: str) -> str:
    """Fetch schema + sample rows from a specific megabase table for summarization."""
    database = MegabaseSessionLocal()
    try:
        init_megabase(database)

        # Fetch schema
        try:
            schema = megabase_describe_table(database, table_name)
        except ValueError:
            return f"Table '{table_name}' not found in the megabase."

        columns = schema.get("columns", [])
        column_names = [col.get("name", "?") for col in columns]
        column_info = "\n".join(
            f"  - {col.get('name', '?')} ({col.get('type', '?')})"
            + (" [PK]" if col.get("primary_key") else "")
            + (" [NOT NULL]" if col.get("nullable") is False else "")
            for col in columns
        )

        lines: list[str] = [
            f"Table: {table_name}",
            f"Columns ({len(column_names)}):",
            column_info,
            "",
        ]

        # Sample rows
        try:
            result = database.execute(sa_text(f'SELECT * FROM "{table_name}" LIMIT {SUMMARY_ROWS_LIMIT}'))
            raw_columns = [str(col) for col in result.keys()]
            rows = result.fetchall()
            lines.append(f"Sample rows ({len(rows)} of up to {SUMMARY_ROWS_LIMIT}):")
            for row in rows[:20]:
                row_dict = dict(zip(raw_columns, row))
                serialized = {k: (str(v) if not isinstance(v, (str, int, float, bool, type(None))) else v) for k, v in row_dict.items()}
                lines.append(json.dumps(serialized, ensure_ascii=True, default=str))
        except Exception as e:
            lines.append(f"Could not read rows: {e}")

        return "\n".join(lines)
    finally:
        database.close()


@router.get("/{group_id}/tables/{table_name}/summarize", response_model=TableSummaryResponse | None)
def get_table_summary(
    group_id: str,
    table_name: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    cached = (
        database.query(LogTableSummary)
        .filter(
            LogTableSummary.group_id == group.id,
            LogTableSummary.table_name == table_name,
        )
        .order_by(LogTableSummary.created_at.desc())
        .first()
    )
    if cached is None:
        return None
    return TableSummaryResponse(**cached.content)


@router.post("/{group_id}/tables/{table_name}/summarize", response_model=TableSummaryResponse)
def summarize_table(
    group_id: str,
    table_name: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    context = _fetch_table_rows_for_summary(str(group.id), table_name)

    system_prompt = (
        "You are an expert log analyst. Analyze the provided log table data and generate a concise, "
        "actionable summary. Return valid JSON matching the required schema. Be concise and factual."
    )

    prompt = (
        f"{context}\n\n"
        "Generate a JSON summary of this log table with these fields:\n"
        '- "summary": a 1-2 sentence overview of what this log is about\n'
        '- "key_observations": a list of key patterns, notable events, or interesting data points (3-6 items)\n'
        '- "severity": one of "low", "medium", "high", "critical" describing the overall severity\n'
        '- "next_actions": a list of actionable next steps or recommended actions (2-4 items)\n'
        '- "errors_or_anomalies": a list of errors, anomalies, or unusual patterns detected (0-3 items)\n'
        "Be concise. Focus on what the data actually shows."
    )

    model = get_generative_model()
    report = model.generate_structured(prompt, TableSummaryResponse, system_prompt=system_prompt)

    # Remove old cached summary for this table, then save the new one
    database.query(LogTableSummary).filter(
        LogTableSummary.group_id == group.id,
        LogTableSummary.table_name == table_name,
    ).delete(synchronize_session="fetch")
    database.flush()

    db_summary = LogTableSummary(
        group_id=group.id,
        table_name=table_name,
        content=report.model_dump(),
    )
    database.add(db_summary)
    database.commit()

    return report


def _get_group_table_names(database: Session, group_id: str) -> set[str]:
    """Return the set of megabase table names registered for a log group."""
    tables = database.query(LogTable).filter(LogTable.group_id == _uuid_or_raw(group_id)).all()
    return {row.table for row in tables}


@router.post("/{group_id}/query", response_model=QueryResponse)
def execute_group_query(
    group_id: str,
    payload: QueryRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    sql_text = payload.sql.strip()

    if not sql_text:
        return QueryResponse(status="error", message="SQL query must not be empty.")

    if FORBIDDEN_SQL_KEYWORDS.search(sql_text):
        return QueryResponse(status="error", message="Only SELECT queries are allowed.")

    parsed = sqlparse.parse(sql_text)
    for statement in parsed:
        if statement.get_type() != "SELECT":
            return QueryResponse(
                status="error",
                message=f"Statement type '{statement.get_type()}' is not allowed. Only SELECT is permitted.",
            )

    allowed_tables = _get_group_table_names(database, str(group.id))
    if not allowed_tables:
        return QueryResponse(status="error", message="No tables are available for this log group.")

    table_error = _validate_table_allowlist(sql_text, allowed_tables)
    if table_error:
        return QueryResponse(status="error", message=table_error)

    sql_text = _inject_sql_limit(sql_text, max_rows=QUERY_RESULT_LIMIT)

    start_time = time.monotonic()
    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        result = megabase_database.execute(sa_text(sql_text))
        raw_columns = list(result.keys()) if result.returns_rows else []
        raw_rows = result.fetchall() if result.returns_rows else []
        elapsed_ms = (time.monotonic() - start_time) * 1000

        columns = [str(col) for col in raw_columns]
        limited_rows = raw_rows

        serializable_rows: list[list[Any]] = []
        for row in limited_rows:
            serializable_row: list[Any] = []
            for value in row:
                if isinstance(value, (dict, list)):
                    serializable_row.append(json.dumps(value, ensure_ascii=True))
                elif hasattr(value, "isoformat"):
                    serializable_row.append(value.isoformat())
                else:
                    serializable_row.append(value)
            serializable_rows.append(serializable_row)

        return QueryResponse(
            status="ok",
            columns=columns,
            rows=serializable_rows,
            row_count=len(raw_rows),
            execution_time_ms=round(elapsed_ms, 2),
            message=f"Returned {len(serializable_rows)} of {len(raw_rows)} rows.",
        )
    except Exception as error:
        elapsed_ms = (time.monotonic() - start_time) * 1000
        logger.exception("SQL query failed for entry %s", group_id)
        return QueryResponse(
            status="error",
            execution_time_ms=round(elapsed_ms, 2),
            message=f"Query failed: {error}",
        )
    finally:
        megabase_database.close()


@router.post("/{group_id}/report")
def generate_group_report(
    group_id: str,
    payload: ReportRequest,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)

    document = Document()

    title_style = document.styles["Title"]
    title_style.font.size = Pt(24)
    document.add_paragraph(payload.title, style="Title")

    document.add_paragraph(f"Log group: {group.name}")
    document.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    document.add_page_break()

    for section in payload.sections:
        document.add_heading(section.heading, level=1)

        for paragraph_text in section.content.split("\n"):
            stripped = paragraph_text.strip()
            if stripped:
                document.add_paragraph(stripped)

        for table_data in section.tables:
            if table_data.title:
                document.add_heading(table_data.title, level=2)

            if not table_data.columns or not table_data.rows:
                continue

            doc_table = document.add_table(rows=1, cols=len(table_data.columns))
            doc_table.style = "Light Grid Accent 1"

            header_cells = doc_table.rows[0].cells
            for index, column_name in enumerate(table_data.columns):
                header_cells[index].text = str(column_name)

            for row_values in table_data.rows:
                row_cells = doc_table.add_row().cells
                for index, value in enumerate(row_values):
                    if index < len(row_cells):
                        row_cells[index].text = str(value) if value is not None else ""

            document.add_paragraph()

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    safe_title = re.sub(r"[^\w\-]", "_", payload.title)[:50] or "report"
    filename = f"{safe_title}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.post("/{group_id}/workbook-report")
def generate_workbook_report(
    group_id: str,
    current_user: User = Depends(get_current_user),
    database: Session = Depends(get_database),
):
    group = _require_owned_group(database=database, group_id=group_id, user_id=current_user.id)
    tables = database.query(LogTable).filter(LogTable.group_id == group.id).all()
    processes = (
        database.query(LogProcess).filter(LogProcess.group_id == group.id).order_by(LogProcess.created_at.desc()).all()
    )

    workbook = Workbook()

    # Summary sheet
    summary_sheet = workbook.active
    summary_sheet.title = "Summary"
    summary_sheet["A1"] = "Logdog Workbook Report"
    summary_sheet["A1"].font = Font(size=16, bold=True)
    summary_sheet["A3"] = "Log Group"
    summary_sheet["B3"] = group.name
    summary_sheet["A4"] = "Generated At"
    summary_sheet["B4"] = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    summary_sheet["A5"] = "Tables"
    summary_sheet["B5"] = len(tables)
    summary_sheet["A6"] = "Processes"
    summary_sheet["B6"] = len(processes)

    process_status_counts: dict[str, int] = {}
    for process in processes:
        process_status_counts[process.status] = process_status_counts.get(process.status, 0) + 1

    summary_sheet["A8"] = "Process Status"
    summary_sheet["B8"] = "Count"
    row_index = 9
    for proc_status, count in process_status_counts.items():
        summary_sheet.cell(row=row_index, column=1, value=proc_status)
        summary_sheet.cell(row=row_index, column=2, value=count)
        row_index += 1

    if process_status_counts:
        chart = BarChart()
        chart.type = "col"
        chart.title = "Process Status"
        chart.y_axis.title = "Count"
        chart.x_axis.title = "Status"
        data = Reference(summary_sheet, min_col=2, min_row=8, max_row=row_index - 1)
        categories = Reference(summary_sheet, min_col=1, min_row=9, max_row=row_index - 1)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(categories)
        chart.shape = 4
        summary_sheet.add_chart(chart, "D3")

    # Data sheets
    megabase_database = MegabaseSessionLocal()
    try:
        init_megabase(megabase_database)
        for table in tables:
            sheet_title = table.table[:31]
            if sheet_title in workbook.sheetnames:
                sheet_title = f"{sheet_title[:28]}_{tables.index(table)}"
            worksheet = workbook.create_sheet(title=sheet_title)

            try:
                result = megabase_database.execute(sa_text(f'SELECT * FROM "{table.table}"'))
                columns = [str(col) for col in result.keys()]
                rows = result.fetchall()

                worksheet.append(columns)
                for row in rows:
                    row_values = []
                    for value in row:
                        if value is None:
                            row_values.append("")
                        elif isinstance(value, (dict, list)):
                            row_values.append(json.dumps(value, ensure_ascii=True))
                        elif isinstance(value, datetime):
                            # Excel does not support timezone-aware datetimes
                            row_values.append(value.replace(tzinfo=None))
                        elif isinstance(value, str) and value.startswith("0000-00-00"):
                            row_values.append("")
                        else:
                            row_values.append(value)
                    worksheet.append(row_values)
            except Exception as error:
                worksheet.append(["Error loading data", str(error)])
    finally:
        megabase_database.close()

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)

    safe_name = re.sub(r"[^\w\-]", "_", group.name)[:50] or "report"
    filename = f"{safe_name}_workbook.xlsx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )
